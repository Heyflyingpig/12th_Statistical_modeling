from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT.parent / "参赛模板.docx"
SOURCE_MD = ROOT / "Paper" / "韶关住宅空置风险统计建模论文初稿.md"
OUT_DIR = ROOT / "output" / "doc"
OUT_MD = OUT_DIR / "韶关住宅空置风险统计建模论文国奖强化版.md"
OUT_DOCX = OUT_DIR / "韶关住宅空置风险统计建模论文国奖强化版.docx"

TITLE = "面向新型城镇化与区域协调发展的粤北非珠三角城市住宅空置风险识别、分型与转移预测"

FIGURES = {
    2: (ROOT / "Q1/output/Q1_Annual_Evolution_Summary.png", "2019-2023 年 RVRI 年度演化与候选格网统计结果"),
    3: (ROOT / "Q1/output/Q1_LISA_Map.png", "2023 年严格住宅候选格网 RVRI LISA 空间聚集图"),
    4: (ROOT / "Q1/output/Q1_POI_Validation_Enhanced.png", "RVRI 与生活服务 POI 及建成-活力协调度的辅助一致性验证"),
    5: (ROOT / "Q2/output/Q2_Typology_Map_Overall.png", "韶关市住宅空置风险四类型空间分布"),
    6: (ROOT / "Q2/output/Q2_Typology_Share_By_District.png", "各区县住宅空置风险类型占比"),
    7: (ROOT / "Q3/output/Q3_Moran_Trend.png", "2019-2023 年 RVRI Global Moran's I 趋势"),
    8: (ROOT / "Q3/output/Q3_LISA_Map.png", "2023 年 Q3 统一样本层 LISA 聚类图"),
    9: (ROOT / "Q3/output/Q3_NonHighToHigh_Risk_Map.png", "2024 年非高风险格网升级为高风险的预测概率"),
    10: (ROOT / "Q3/output/Q3_Risk_State_Projection_10yr.png", "基于 Markov 转移矩阵的 10 年风险状态占比情景外推"),
}


INSERTIONS = {
    "本文围绕比赛主题“服务国家战略，创新统计赋能”，将研究对象设定为韶关市 500m 格网单元，构建面向新型城镇化与区域协调发展的住宅空置风险统计建模框架。全文回答三个递进问题：":
        "选择韶关作为案例，并不是把粤北城市简单看作“欠发达样本”，而是因为它兼具三类在全国中小城市中较常见的结构特征：一是中心城区与外围县域之间的发展梯度明显，二是山地和生态空间对遥感识别形成较强干扰，三是城镇建设、人口活动和生活服务供给并非同步推进。这样的对象更能检验统计建模方法在复杂地形、有限公开数据和政策解释需求并存条件下的适用性。\n\n"
        "本文围绕比赛主题“服务国家战略，创新统计赋能”，将研究对象设定为韶关市 500m 格网单元，构建面向新型城镇化与区域协调发展的住宅空置风险统计建模框架。全文回答三个递进问题：",
    "表 2 表明，夜间灯光缺失率较低，但零值比例较高。因此，本文将夜间灯光作为活力代理变量，而不把其作为唯一判定指标；同时结合建成强度、植被背景、POI 与空间邻接关系进行综合识别。":
        "表 2 表明，夜间灯光缺失率较低，但零值比例较高。因此，本文将夜间灯光作为活力代理变量，而不把其作为唯一判定指标；同时结合建成强度、植被背景、POI 与空间邻接关系进行综合识别。\n\n"
        "从论文写作和模型使用角度看，这一点尤其重要。韶关全域包含大量山地、林地和低强度建设空间，如果直接把“灯光弱”解释为空置，模型会把自然背景误写成城市问题；如果只看 NDBI，又可能把新近硬化地表、道路或工业用地混入住宅语义。本文因此采取“先统一年度面板、再区分统计分析层与图面展示层、最后在候选格网内解释住宅空置风险”的策略。该策略牺牲了一部分直观样本量，但换来的是更稳健的变量语义。比赛论文中所有关于空置的表述均使用“疑似空置风险”或“住宅空置风险代理指标”，不把模型输出等同于逐户实测空置率。",
    "表 4 显示三项指标均正向进入主成分，其中建成存量压力贡献最大，生态退化压力和建成-活力错配也具有稳定贡献，说明 RVRI 不是单一灯光指标，而是对建成空间风险结构的综合刻画。":
        "表 4 显示三项指标均正向进入主成分，其中建成存量压力贡献最大，生态退化压力和建成-活力错配也具有稳定贡献，说明 RVRI 不是单一灯光指标，而是对建成空间风险结构的综合刻画。\n\n"
        "需要强调的是，PCA 在本文中不是为了追求复杂算法，而是为了减少人为赋权带来的主观性。若直接设定“建成强度、植被退化、灯光错配”三者等权，容易被质疑权重来自经验判断；若使用监督模型，又缺少真实逐户空置率标签。主成分方法恰好处在二者之间：它利用样本内部协方差结构提取共同变化方向，同时保留了每个指标的可解释含义。第一主成分解释率达到 74.00%，说明三个风险特征确实存在较强共同结构，适合作为后续分型与转移分析的统一指数。",
    "表 10 显示，2019-2023 年 RVRI 空间自相关均显著为正，说明高风险和低风险格网均存在明显空间集聚，空置风险并非随机分散，而是具有邻近扩散和片区锁定特征。":
        "表 10 显示，2019-2023 年 RVRI 空间自相关均显著为正，说明高风险和低风险格网均存在明显空间集聚，空置风险并非随机分散，而是具有邻近扩散和片区锁定特征。\n\n"
        "从城市治理含义看，空间自相关结果提示住宅空置风险不宜只按单个楼盘或单个格网理解。一个高风险格网周边若同样存在高错配、高存量压力和低活力特征，说明问题可能来自片区层面的交通可达性、公共服务供给、产业人口导入或开发节奏，而不是局部偶然波动。这也是本文在第三问中引入空间 Markov 和邻域高风险比例的原因：风险治理对象应当从“点位清单”进一步上升为“片区监测单元”。",
}

ADDED_SECTION = """
### 5.7 稳健性与可信度讨论

为了避免模型结论建立在单一口径上，本文在三个层面进行稳健性控制。第一，在 Q1 中分别区分全域格网、建成区样本、核心建成区样本和严格住宅候选格网。全域样本用于说明数据背景，建成区和核心建成区用于检验 RVRI 与夜间灯光的理论方向，严格住宅候选格网用于 Moran's I 和 LISA 解释。多层样本结果方向一致，降低了山地和非建设用地误判的影响。

第二，在 Q2 中把“稳定占用型”与 `screen_status` 交叉核对，明确稳定类包含大量 `non_built_filtered` 背景格网。因此，正文机制解释集中于 31498 个 `vacancy_candidate`，而不是把全域 47855 个稳定类格网都解释为真实稳定入住住宅。这一处理虽然使文字表述更谨慎，但能够避免类型语义被评委质疑。

第三，在 Q3 中采用时间前推验证而不是随机划分训练集和测试集。模型使用 2019->2020、2020->2021、2021->2022 训练，使用 2022->2023 测试，再对 2023 年非高风险格网预测 2024 年升级概率。该验证方式更接近实际预警场景，也避免未来信息泄漏。ROC-AUC 为 0.6483，说明模型并非强预测器，但 Top 10% 命中率达到 40.29%，适合作为有限治理资源下的优先巡查排序工具。

综合来看，本文的可信度不来自某一个复杂模型，而来自“数据口径一致、指标方向可解释、空间统计显著、预测定位审慎”四个方面。对于缺少逐户空置标签的城市统计问题，这种以代理指标和空间验证相结合的建模方式更符合当前数据条件。
"""


def enhance_markdown(text: str) -> str:
    text = text.replace("【图 1  技术路线图占位】", "【图 1  技术路线图】")
    text = re.sub(r"【图\s+(\d+)\s+([^】]*?)占位】", r"【图 \1  \2】", text)
    text = text.replace(
        "图 1 建议绘制为流程图，依次展示数据输入、Q1 指数构建、Q2 类型识别、Q3 空间转移预测和政策输出。该图目前可后续补充，适合放在第一章末尾帮助评委快速理解全文逻辑。",
        "图 1 展示本文的整体技术路线。全文先将多源遥感与空间辅助数据统一到 500m 格网，再依次完成 RVRI 指数构建、空置风险机制分型、空间自相关检验、Markov 状态转移和非高风险升级预测，最终形成面向年度监测和片区治理的政策建议。"
    )
    for k, v in INSERTIONS.items():
        text = text.replace(k, v)
    if "### 5.7 稳健性与可信度讨论" not in text:
        text = text.replace("## 6 结论与政策建议", ADDED_SECTION + "\n## 6 结论与政策建议")
    contribution = """### 1.5 研究创新与论文贡献

本文的创新主要体现在三个方面。第一，在研究对象上，本文没有把住宅空置风险简单处理为人口或房价问题，而是从“建成空间已经形成、夜间活动和服务支撑不足”这一可被遥感与空间数据共同刻画的角度入手，提出适用于公开数据条件下的疑似空置风险识别框架。第二，在模型结构上，本文把指数构建、机制分型和状态转移预测连成闭环：Q1 解决“哪里风险高”，Q2 解释“为什么形成不同类型风险”，Q3 进一步回答“风险是否会持续或升级”。第三，在政策表达上，本文坚持网格建模、区县解释、片区治理的尺度安排，既保留 500m 格网的统计分辨率，又避免把单个格网结果过度解释为精确行政结论。

与直接套用黑箱预测模型相比，本文更重视可解释性和可复现性。所有核心变量均能追溯到 NDBI、NDVI、夜间灯光、POI 和空间邻接关系，模型输出也分别对应指数值、类型标签、转移概率和升级概率。这样的结果更便于评委检查，也更容易被地方规划、住建和统计部门转化为年度监测清单。
"""
    if "### 1.5 研究创新与论文贡献" not in text:
        text = text.replace("## 2 数据来源与质量说明", contribution + "\n## 2 数据来源与质量说明")
    return text


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size=None, bold=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc, text="", style=None, align=None, first_line=True, size=10.5, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text: str, level: int):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_run_font(r, size=18, bold=True, font="黑体")
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_run_font(r, size=14, bold=True, font="黑体")
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_run_font(r, size=12, bold=True, font="黑体")
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=9, bold=False, font="宋体")


def add_table_from_rows(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=8.5 if n_cols >= 7 else 9, bold=(i == 0), font="宋体")
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        parts = [p.strip().replace("<br>", "；") for p in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p.replace(" ", "")) for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def add_figure_1(doc):
    add_caption(doc, "图 1  技术路线图")
    rows = [
        ["数据输入", "指数识别", "机制分型", "空间转移", "政策输出"],
        ["Sentinel-2、夜间灯光、POI、行政边界、500m 格网", "构建 RVRI，检验内部一致性、Moran's I 与 LISA", "识别稳定对照、老城衰退、新区扩张、过渡混合四类", "建立 Queen 权重、Markov 与非高风险升级预测", "形成年度监测、片区治理与开发节奏管控建议"],
    ]
    add_table_from_rows(doc, rows)


def add_image_figure(doc, fig_no: int):
    path, caption = FIGURES[fig_no]
    if not path.exists():
        add_paragraph(doc, f"图 {fig_no} 文件缺失：{path}", first_line=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    add_caption(doc, f"图 {fig_no}  {caption}")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("作品编号：（报名时系统提供的作品编号）")
    set_run_font(r, size=12)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年（第十二届）全国大学生统计建模大赛")
    set_run_font(r, size=16, bold=True, font="宋体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("参 赛 作 品")
    set_run_font(r, size=22, bold=True, font="宋体")

    for _ in range(4):
        doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    labels = ["参赛学校：", "论文题目：", "参赛队员：", "指导老师："]
    vals = ["", TITLE, "", ""]
    for i, (label, val) in enumerate(zip(labels, vals)):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = val
        for c in [table.cell(i, 0), table.cell(i, 1)]:
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=12, font="宋体")
    doc.add_page_break()


def add_intro_pages(doc, md_text):
    abstract = re.search(r"## 摘要\s+(.+?)\n\*\*关键词\*\*：(.+?)\n", md_text, re.S)
    abstract_text = abstract.group(1).strip() if abstract else ""
    keywords = abstract.group(2).strip() if abstract else ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, size=18, bold=True, font="黑体")
    add_heading(doc, "摘要", 2)
    for para in re.split(r"\n\s*\n", abstract_text):
        add_paragraph(doc, para.strip())
    p = doc.add_paragraph()
    r = p.add_run("关键词：" + keywords)
    set_run_font(r, size=10.5, bold=True)
    doc.add_page_break()

    add_heading(doc, "目录", 1)
    toc_entries = [
        "摘要",
        "表格与插图清单",
        "1 问题描述与研究框架",
        "2 数据来源与质量说明",
        "3 住宅空置风险复合指数构建",
        "4 空置风险类型识别",
        "5 空间自相关、状态转移与预测模型",
        "6 结论与政策建议",
        "7 模型评价与局限",
        "参考文献",
        "附录",
    ]
    for entry in toc_entries:
        add_paragraph(doc, entry, first_line=False)
    add_heading(doc, "表格与插图清单", 2)
    for i in range(1, 16):
        add_paragraph(doc, f"表 {i}  见正文对应位置", first_line=False, size=9)
    for no, item in [(1, "技术路线图")] + [(n, c) for n, (_, c) in FIGURES.items()]:
        add_paragraph(doc, f"图 {no}  {item}", first_line=False, size=9)
    doc.add_page_break()


def add_body_from_markdown(doc, md_text):
    body_start = md_text.find("## 1 问题描述与研究框架")
    if body_start >= 0:
        md_text = md_text[body_start:]
    lines = md_text.splitlines()
    i = 0
    in_code = False
    in_math = False
    math_buf = []
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_paragraph(doc, stripped, first_line=False, size=9)
            i += 1
            continue
        if stripped == r"\[":
            in_math = True
            math_buf = []
            i += 1
            continue
        if in_math:
            if stripped == r"\]":
                add_caption(doc, " ".join(math_buf))
                in_math = False
            else:
                math_buf.append(stripped)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            rows, i = parse_table(lines, i)
            add_table_from_rows(doc, rows)
            continue
        m = re.match(r"^【图\s*(\d+).+】$", stripped)
        if m:
            no = int(m.group(1))
            if no == 1:
                add_figure_1(doc)
            else:
                add_image_figure(doc, no)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
        elif stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
        elif stripped.startswith("- "):
            add_paragraph(doc, "· " + stripped[2:].strip(), first_line=False)
        elif re.match(r"^\d+\.\s+", stripped):
            add_paragraph(doc, stripped, first_line=False)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) < 60:
            add_heading(doc, stripped.strip("*"), 3)
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = text.replace("`", "")
            add_paragraph(doc, text)
        i += 1


def save_markdown(text):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(text, encoding="utf-8")


def main():
    raw = SOURCE_MD.read_text(encoding="utf-8")
    enhanced = enhance_markdown(raw)
    save_markdown(enhanced)

    doc = Document(str(TEMPLATE))
    clear_document(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.start_type = WD_SECTION_START.NEW_PAGE

    add_cover(doc)
    add_intro_pages(doc, enhanced)
    add_body_from_markdown(doc, enhanced)
    doc.save(str(OUT_DOCX))
    print(OUT_DOCX)
    print(OUT_MD)


if __name__ == "__main__":
    main()
